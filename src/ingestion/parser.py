"""
NovelForge 文档解析器
支持DOCX/MD/TXT格式的文档解析和智能分块
"""

import re
import logging
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """文档类型"""
    WORLD = "world"          # 世界观设定
    CHARACTER = "character"  # 角色资料
    STYLE = "style"          # 写作风格
    REFERENCE = "reference"  # 参考资料
    CHAPTER = "chapter"      # 已有正文
    OTHER = "other"


@dataclass
class DocumentChunk:
    """文档分块"""
    id: str
    document_id: str
    chunk_index: int
    content: str
    metadata: Optional[Dict] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class ParsedDocument:
    """解析后的文档"""
    id: str
    name: str
    doc_type: DocumentType
    content: str
    word_count: int
    chunks: List[DocumentChunk]
    metadata: Optional[Dict] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.supported_formats = {'.txt', '.md', '.docx'}
    
    def parse(self, file_path: str, doc_type: DocumentType = DocumentType.OTHER,
              chunk_size: int = 1000, chunk_overlap: int = 200) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file_path: 文件路径
            doc_type: 文档类型
            chunk_size: 分块大小(字符数)
            chunk_overlap: 分块重叠(字符数)
            
        Returns:
            解析后的文档
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {suffix}")
        
        # 读取内容
        content = self._read_file(path, suffix)
        
        # 分块
        chunks = self._chunk_text(content, chunk_size, chunk_overlap)
        
        # 创建文档对象
        doc_id = self._generate_id(path)
        document = ParsedDocument(
            id=doc_id,
            name=path.name,
            doc_type=doc_type,
            content=content,
            word_count=len(content),
            chunks=[],
            metadata={
                "file_path": str(path),
                "file_size": path.stat().st_size,
                "format": suffix,
            }
        )
        
        # 创建分块对象
        for i, chunk_text in enumerate(chunks):
            chunk = DocumentChunk(
                id=f"{doc_id}_chunk_{i}",
                document_id=doc_id,
                chunk_index=i,
                content=chunk_text,
                metadata={
                    "start_char": sum(len(c) for c in chunks[:i]),
                    "end_char": sum(len(c) for c in chunks[:i+1]),
                }
            )
            document.chunks.append(chunk)
        
        logger.info(f"文档解析完成: {path.name}, {len(chunks)}个分块")
        return document
    
    def _read_file(self, path: Path, suffix: str) -> str:
        """读取文件内容"""
        if suffix in ('.txt', '.md'):
            return self._read_text_file(path)
        elif suffix == '.docx':
            return self._read_docx_file(path)
        else:
            raise ValueError(f"不支持的格式: {suffix}")
    
    def _read_text_file(self, path: Path) -> str:
        """读取文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"无法解码文件: {path}")
    
    def _read_docx_file(self, path: Path) -> str:
        """读取DOCX文件"""
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            return '\n\n'.join(paragraphs)
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
    
    def _chunk_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """
        智能分块
        
        优先按段落分割，其次按句子分割，最后按字符分割
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        
        # 先按段落分割
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        
        for para in paragraphs:
            # 如果当前块加上新段落不超过限制
            if len(current_chunk) + len(para) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += '\n\n'
                current_chunk += para
            else:
                # 保存当前块
                if current_chunk:
                    chunks.append(current_chunk)
                
                # 如果单个段落超过限制，需要进一步分割
                if len(para) > chunk_size:
                    sub_chunks = self._chunk_by_sentences(para, chunk_size, chunk_overlap)
                    chunks.extend(sub_chunks)
                    current_chunk = ""
                else:
                    current_chunk = para
        
        # 保存最后一个块
        if current_chunk:
            chunks.append(current_chunk)
        
        # 添加重叠
        if chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks, chunk_overlap)
        
        return chunks
    
    def _chunk_by_sentences(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """按句子分割"""
        # 中文句子分割
        sentences = re.split(r'([。！？；\n])', text)
        
        chunks = []
        current_chunk = ""
        
        for i in range(0, len(sentences), 2):
            sentence = sentences[i]
            # 添加标点符号
            if i + 1 < len(sentences):
                sentence += sentences[i + 1]
            
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _add_overlap(self, chunks: List[str], overlap_size: int) -> List[str]:
        """添加重叠"""
        result = [chunks[0]]
        
        for i in range(1, len(chunks)):
            # 从上一个块的末尾取overlap_size个字符
            prev_chunk = chunks[i - 1]
            overlap = prev_chunk[-overlap_size:] if len(prev_chunk) > overlap_size else prev_chunk
            
            # 将重叠部分添加到当前块的开头
            result.append(overlap + chunks[i])
        
        return result
    
    def _generate_id(self, path: Path) -> str:
        """生成文档ID"""
        import uuid
        return uuid.uuid4().hex[:12]


class DocumentClassifier:
    """文档分类器"""
    
    # 关键词映射
    KEYWORD_MAP = {
        DocumentType.WORLD: ['世界观', '设定', '背景', '规则', '历史', '地理', '国家', '种族'],
        DocumentType.CHARACTER: ['角色', '人物', '主角', '配角', '性格', '背景', '能力', '关系'],
        DocumentType.STYLE: ['风格', '写法', '技法', '要求', '禁忌', '语言', '文风'],
        DocumentType.CHAPTER: ['第章', '章节', '正文', '内容'],
    }
    
    @classmethod
    def classify(cls, content: str, filename: str = "") -> DocumentType:
        """
        自动分类文档类型
        
        Args:
            content: 文档内容
            filename: 文件名
            
        Returns:
            文档类型
        """
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        # 统计关键词出现次数
        scores = {}
        
        for doc_type, keywords in cls.KEYWORD_MAP.items():
            score = 0
            for keyword in keywords:
                if keyword in content_lower or keyword in filename_lower:
                    score += 1
            scores[doc_type] = score
        
        # 返回得分最高的类型
        if scores:
            max_type = max(scores, key=lambda document_type: scores[document_type])
            if scores[max_type] > 0:
                return max_type
        
        return DocumentType.OTHER


class TextCleaner:
    """文本清理器"""
    
    @staticmethod
    def clean(text: str) -> str:
        """清理文本"""
        # 移除多余空白
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # 移除特殊字符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        
        # 标准化引号
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        
        return text.strip()
    
    @staticmethod
    def extract_metadata(text: str) -> Dict:
        """提取元数据"""
        metadata = {}
        
        # 提取标题
        title_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
        if title_match:
            metadata['title'] = title_match.group(1)
        
        # 统计字数
        metadata['char_count'] = len(text)
        metadata['word_count'] = len(text.split())
        
        # 统计段落数
        metadata['paragraph_count'] = len([p for p in text.split('\n\n') if p.strip()])
        
        return metadata


# 便捷函数
def parse_document(file_path: str, doc_type: str = "auto",
                   chunk_size: int = 1000, chunk_overlap: int = 200) -> ParsedDocument:
    """
    解析文档的便捷函数
    
    Args:
        file_path: 文件路径
        doc_type: 文档类型 ("auto", "world", "character", "style", "reference", "chapter")
        chunk_size: 分块大小
        chunk_overlap: 分块重叠
        
    Returns:
        解析后的文档
    """
    parser = DocumentParser()
    
    # 自动分类
    if doc_type == "auto":
        content = parser._read_file(Path(file_path), Path(file_path).suffix)
        doc_type = DocumentClassifier.classify(content, Path(file_path).name)
    else:
        doc_type = DocumentType(doc_type)
    
    return parser.parse(file_path, doc_type, chunk_size, chunk_overlap)


def batch_parse(file_paths: List[str], doc_type: str = "auto",
                chunk_size: int = 1000) -> List[ParsedDocument]:
    """
    批量解析文档
    
    Args:
        file_paths: 文件路径列表
        doc_type: 文档类型
        chunk_size: 分块大小
        
    Returns:
        解析后的文档列表
    """
    parser = DocumentParser()
    documents = []
    
    for file_path in file_paths:
        try:
            doc = parse_document(file_path, doc_type, chunk_size)
            documents.append(doc)
        except Exception as e:
            logger.error(f"解析文档失败 {file_path}: {e}")
    
    return documents
