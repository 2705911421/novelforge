"""SQLite-authoritative export service with history tracking."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional, cast

from src.core.database import Database, generate_id


def _sanitize_filename(name: str) -> str:
    """清洗文件名，防止路径穿越 (BLOCKING fix)

    Args:
        name: 原始文件名

    Returns:
        安全的文件名
    """
    # 移除路径分隔符和特殊字符
    sanitized = re.sub(r'[/\\:*?"<>|..]', '_', name)
    # 移除前后空白和点号
    sanitized = sanitized.strip('. ')
    # 如果为空，使用默认名称
    return sanitized or "untitled"


class ExportService:
    """Export books from SQLite to various formats with history tracking."""

    def __init__(self, db: Database, output_dir: Path):
        self.db = db
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_book(
        self,
        project_id: str,
        book_id: str,
        format: str = "md",
        approved_only: bool = False,
    ) -> dict[str, Any]:
        """Export a book from SQLite to a file.
        
        Args:
            project_id: The project ID
            book_id: The book ID
            format: Export format (md, txt, docx)
            approved_only: Only export approved chapters
            
        Returns:
            Export record with file_path, word_count, chapter_count
        """
        # Get book info.
        book = self.db.fetchone(
            "SELECT * FROM books WHERE id=? AND project_id=?",
            (book_id, project_id),
        )
        if not book:
            raise ValueError(f"Book not found: {book_id}")

        # Get chapters.
        if approved_only:
            chapters = self.db.fetchall(
                """SELECT * FROM chapters 
                   WHERE book_id=? AND status IN ('approved', 'committed', 'exported')
                   ORDER BY number""",
                (book_id,),
            )
        else:
            chapters = self.db.fetchall(
                """SELECT * FROM chapters 
                   WHERE book_id=? AND status != 'planned'
                   ORDER BY number""",
                (book_id,),
            )

        if not chapters:
            raise ValueError("No chapters to export")

        # Generate output file.
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{_sanitize_filename(book['title'])}_{timestamp}.{format}"
        file_path = self.output_dir / filename

        # Export based on format.
        if format == "md":
            content = self._export_markdown(book, chapters)
        elif format == "txt":
            content = self._export_txt(book, chapters)
        elif format == "docx":
            self._export_docx(book, chapters, file_path)
            content = None
        else:
            raise ValueError(f"Unsupported format: {format}")

        # Write file.
        if content is not None:
            file_path.write_text(content, encoding="utf-8")

        # Calculate stats.
        word_count = sum(len(ch.get("content", "")) for ch in chapters)

        # Record export in database.
        export_id = generate_id()
        now = datetime.now().isoformat()
        with self.db.transaction() as conn:
            conn.execute(
                """INSERT INTO exports(id, project_id, book_id, format, file_path, 
                   file_size, chapter_count, word_count, approved_only, status, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    export_id, project_id, book_id, format, str(file_path),
                    file_path.stat().st_size, len(chapters), word_count,
                    approved_only, "completed", now,
                ),
            )

        return {
            "id": export_id,
            "file_path": str(file_path),
            "format": format,
            "chapter_count": len(chapters),
            "word_count": word_count,
            "file_size": file_path.stat().st_size,
        }

    def _export_docx(self, book: dict, chapters: list[dict], path: Path) -> None:
        """Write a real Word document using the same SQLite read model."""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.styles.style import ParagraphStyle
            from docx.shared import Pt
        except ImportError as exc:
            raise ValueError("导出 Word 需要安装 python-docx") from exc

        document = Document()
        normal = cast(ParagraphStyle, document.styles["Normal"])
        normal.font.name = "宋体"
        normal.font.size = Pt(12)
        title = document.add_heading(str(book.get("title") or "未命名作品"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"类型: {book.get('genre', '未分类')}")
        document.add_paragraph(f"总字数: {sum(len(ch.get('content', '')) for ch in chapters)}")
        document.add_paragraph(f"章节数: {len(chapters)}")
        document.add_page_break()
        for index, chapter in enumerate(chapters):
            document.add_heading(
                f"第{chapter['number']}章 {chapter.get('title') or '未命名章节'}",
                level=1,
            )
            for paragraph in str(chapter.get("content") or "").splitlines():
                text = paragraph.strip()
                if text:
                    document.add_paragraph(text)
            if index < len(chapters) - 1:
                document.add_page_break()
        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(path))

    def get_export_history(self, project_id: str) -> list[dict[str, Any]]:
        """Get export history for a project."""
        return self.db.fetchall(
            """SELECT * FROM exports 
               WHERE project_id=?
               ORDER BY created_at DESC""",
            (project_id,),
        )

    def get_export(self, export_id: str) -> Optional[dict[str, Any]]:
        """Get a specific export record."""
        return self.db.fetchone(
            "SELECT * FROM exports WHERE id=?",
            (export_id,),
        )

    def _export_markdown(self, book: dict, chapters: list[dict]) -> str:
        """Export to Markdown format."""
        lines = [
            f"# {book['title']}",
            "",
            f"**类型**: {book.get('genre', '未分类')}",
            f"**总字数**: {sum(len(ch.get('content', '')) for ch in chapters)}",
            f"**章节数**: {len(chapters)}",
            "",
            "---",
            "",
        ]

        for chapter in chapters:
            lines.append(f"## 第{chapter['number']}章 {chapter.get('title', '')}")
            lines.append("")
            lines.append(chapter.get("content", ""))
            lines.append("")
            lines.append("---")
            lines.append("")

        return "\n".join(lines)

    def _export_txt(self, book: dict, chapters: list[dict]) -> str:
        """Export to plain text format."""
        lines = [
            book['title'],
            "=" * len(book['title']) * 2,
            "",
            f"类型: {book.get('genre', '未分类')}",
            f"总字数: {sum(len(ch.get('content', '')) for ch in chapters)}",
            f"章节数: {len(chapters)}",
            "",
            "-" * 40,
            "",
        ]

        for chapter in chapters:
            lines.append(f"第{chapter['number']}章 {chapter.get('title', '')}")
            lines.append("")
            lines.append(chapter.get("content", ""))
            lines.append("")
            lines.append("-" * 40)
            lines.append("")

        return "\n".join(lines)

    # ========== EXPORT-004: Story Bible 导出 ==========

    def export_story_bible(
        self,
        project_id: str,
        book_id: str,
        format: str = "md",
    ) -> dict[str, Any]:
        """导出 Story Bible (EXPORT-004)

        Args:
            project_id: 项目ID
            book_id: 书籍ID
            format: 导出格式 (md/txt)

        Returns:
            导出记录
        """
        # 获取 Story Bible 数据
        bible_data = self._get_story_bible_data(project_id, book_id)

        if not bible_data:
            raise ValueError("没有 Story Bible 数据可导出")

        # 生成输出文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        book = self.db.fetchone("SELECT title FROM books WHERE id=?", (book_id,))
        book_title = book["title"] if book else "unknown"
        filename = f"{_sanitize_filename(book_title)}_story_bible_{timestamp}.{format}"
        file_path = self.output_dir / filename

        # 根据格式导出
        if format == "md":
            content = self._export_story_bible_markdown(bible_data)
        elif format == "txt":
            content = self._export_story_bible_txt(bible_data)
        else:
            raise ValueError(f"不支持的格式: {format}")

        # 写入文件
        file_path.write_text(content, encoding="utf-8")

        # 记录导出
        export_id = generate_id()
        now = datetime.now().isoformat()
        with self.db.transaction() as conn:
            conn.execute(
                """INSERT INTO exports(id, project_id, book_id, format, file_path,
                   file_size, status, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (export_id, project_id, book_id, format, str(file_path),
                 file_path.stat().st_size, "completed", now),
            )

        return {
            "id": export_id,
            "file_path": str(file_path),
            "format": format,
            "file_size": file_path.stat().st_size,
        }

    def _get_story_bible_data(self, project_id: str, book_id: str) -> dict[str, Any]:
        """获取 Story Bible 数据"""
        # 从 story_bible_snapshots 获取最新快照
        snapshot = self.db.fetchone(
            """SELECT * FROM story_bible_snapshots
               WHERE workspace_id IN (
                   SELECT id FROM story_bible_workspaces WHERE project_id = ?
               )
               ORDER BY created_at DESC LIMIT 1""",
            (project_id,),
        )

        if snapshot:
            import json
            return json.loads(snapshot.get("payload", "{}"))

        # 从 story_bible_steps 获取已确认的步骤
        steps = self.db.fetchall(
            """SELECT sbs.* FROM story_bible_steps sbs
               JOIN story_bible_workspaces sbw ON sbw.id = sbs.workspace_id
               WHERE sbw.project_id = ? AND sbs.status = 'confirmed'
               ORDER BY sbs.step_number""",
            (project_id,),
        )

        if steps:
            import json
            result = {}
            for step in steps:
                key = step.get("step_key", "")
                draft = step.get("draft", "{}")
                if isinstance(draft, str):
                    try:
                        draft_data = json.loads(draft)
                        result[key] = draft_data.get("content", draft)
                    except (json.JSONDecodeError, TypeError):
                        result[key] = draft
                elif isinstance(draft, dict):
                    result[key] = draft.get("content", str(draft))
            return result

        return {}

    def _export_story_bible_markdown(self, data: dict[str, Any]) -> str:
        """导出 Story Bible 为 Markdown"""
        lines = [
            "# Story Bible",
            "",
            "---",
            "",
        ]

        for key, value in data.items():
            if isinstance(value, str) and value:
                lines.append(f"## {key}")
                lines.append("")
                lines.append(value)
                lines.append("")
            elif isinstance(value, dict):
                lines.append(f"## {key}")
                lines.append("")
                for sub_key, sub_value in value.items():
                    if isinstance(sub_value, str) and sub_value:
                        lines.append(f"### {sub_key}")
                        lines.append("")
                        lines.append(sub_value)
                        lines.append("")

        return "\n".join(lines)

    def _export_story_bible_txt(self, data: dict[str, Any]) -> str:
        """导出 Story Bible 为纯文本"""
        lines = [
            "Story Bible",
            "=" * 20,
            "",
        ]

        for key, value in data.items():
            if isinstance(value, str) and value:
                lines.append(f"[{key}]")
                lines.append(value)
                lines.append("")
            elif isinstance(value, dict):
                lines.append(f"[{key}]")
                for sub_key, sub_value in value.items():
                    if isinstance(sub_value, str) and sub_value:
                        lines.append(f"  {sub_key}: {sub_value}")
                lines.append("")

        return "\n".join(lines)

    # ========== EXPORT-005: 审查报告导出 ==========

    def export_review_report(
        self,
        project_id: str,
        book_id: str,
        format: str = "md",
    ) -> dict[str, Any]:
        """导出审查报告 (EXPORT-005)

        Args:
            project_id: 项目ID
            book_id: 书籍ID
            format: 导出格式 (md/txt)

        Returns:
            导出记录
        """
        # 获取审查数据
        reviews = self._get_review_data(book_id)

        if not reviews:
            raise ValueError("没有审查数据可导出")

        # 生成输出文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        book = self.db.fetchone("SELECT title FROM books WHERE id=?", (book_id,))
        book_title = book["title"] if book else "unknown"
        filename = f"{_sanitize_filename(book_title)}_review_report_{timestamp}.{format}"
        file_path = self.output_dir / filename

        # 根据格式导出
        if format == "md":
            content = self._export_review_report_markdown(reviews)
        elif format == "txt":
            content = self._export_review_report_txt(reviews)
        else:
            raise ValueError(f"不支持的格式: {format}")

        # 写入文件
        file_path.write_text(content, encoding="utf-8")

        # 记录导出
        export_id = generate_id()
        now = datetime.now().isoformat()
        with self.db.transaction() as conn:
            conn.execute(
                """INSERT INTO exports(id, project_id, book_id, format, file_path,
                   file_size, status, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (export_id, project_id, book_id, format, str(file_path),
                 file_path.stat().st_size, "completed", now),
            )

        return {
            "id": export_id,
            "file_path": str(file_path),
            "format": format,
            "file_size": file_path.stat().st_size,
        }

    def _get_review_data(self, book_id: str) -> list[dict[str, Any]]:
        """获取审查数据"""
        # 从 reviews 表获取审查信息
        reviews = self.db.fetchall(
            """SELECT r.*, c.number, c.title
               FROM reviews r
               JOIN chapters c ON c.id = r.chapter_id
               WHERE c.book_id = ?
               ORDER BY c.number, r.created_at DESC""",
            (book_id,),
        )

        result = []
        for review in reviews:
            review_dict = dict(review)

            # 获取维度评分
            dimensions = self.db.fetchall(
                """SELECT * FROM review_dimensions WHERE review_id = ?""",
                (review["id"],),
            )
            review_dict["dimensions_list"] = [dict(d) for d in dimensions]

            # 获取问题列表
            issues = self.db.fetchall(
                """SELECT * FROM review_issues WHERE review_id = ?""",
                (review["id"],),
            )
            review_dict["issues_list"] = [dict(i) for i in issues]

            result.append(review_dict)

        return result

    def _export_review_report_markdown(self, reviews: list[dict]) -> str:
        """导出审查报告为 Markdown"""
        lines = [
            "# 审查报告",
            "",
            "---",
            "",
        ]

        total_score = 0
        review_count = 0
        for review in reviews:
            score = review.get("overall_score", 0) or 0
            total_score += score
            review_count += 1

            lines.append(f"## 第{review['number']}章 {review.get('title', '')}")
            lines.append("")
            lines.append(f"**评分**: {score:.1f}")
            lines.append(f"**结论**: {review.get('verdict', 'N/A')}")
            lines.append(f"**通过**: {'是' if review.get('passed') else '否'}")
            lines.append("")

            # 维度详情
            dimensions = review.get("dimensions_list", [])
            if dimensions:
                lines.append("### 维度评分")
                lines.append("")
                for dim in dimensions:
                    dim_name = dim.get("dimension", "")
                    dim_score = dim.get("score", 0)
                    if dim_name:
                        lines.append(f"- {dim_name}: {dim_score:.1f}")
                lines.append("")

            # 问题列表
            issues = review.get("issues_list", [])
            if issues:
                lines.append("### 问题")
                lines.append("")
                for issue in issues:
                    severity = issue.get("severity", "")
                    description = issue.get("description", "")
                    lines.append(f"- [{severity}] {description}")
                lines.append("")

            lines.append("---")
            lines.append("")

        # 汇总
        if review_count > 0:
            avg_score = total_score / review_count
            lines.insert(3, f"**平均评分**: {avg_score:.1f}")
            lines.insert(4, f"**审查次数**: {review_count}")
            lines.insert(5, "")

        return "\n".join(lines)

    def _export_review_report_txt(self, reviews: list[dict]) -> str:
        """导出审查报告为纯文本"""
        lines = [
            "审查报告",
            "=" * 20,
            "",
        ]

        total_score = 0
        review_count = 0
        for review in reviews:
            score = review.get("overall_score", 0) or 0
            total_score += score
            review_count += 1

            lines.append(f"第{review['number']}章 {review.get('title', '')}")
            lines.append(f"评分: {score:.1f}")
            lines.append(f"结论: {review.get('verdict', 'N/A')}")
            lines.append(f"通过: {'是' if review.get('passed') else '否'}")

            # 维度详情
            dimensions = review.get("dimensions_list", [])
            if dimensions:
                lines.append("维度评分:")
                for dim in dimensions:
                    dim_name = dim.get("dimension", "")
                    dim_score = dim.get("score", 0)
                    if dim_name:
                        lines.append(f"  {dim_name}: {dim_score:.1f}")

            # 问题列表
            issues = review.get("issues_list", [])
            if issues:
                lines.append("问题:")
                for issue in issues:
                    severity = issue.get("severity", "")
                    description = issue.get("description", "")
                    lines.append(f"  [{severity}] {description}")

            lines.append("-" * 40)
            lines.append("")

        # 汇总
        if review_count > 0:
            avg_score = total_score / review_count
            lines.insert(3, f"平均评分: {avg_score:.1f}")
            lines.insert(4, f"审查次数: {review_count}")
            lines.insert(5, "")

        return "\n".join(lines)

    # ========== EXPORT-006: 伏笔表导出 ==========

    def export_foreshadowing(
        self,
        project_id: str,
        book_id: str,
        format: str = "md",
        status_filter: Optional[str] = None,
    ) -> dict[str, Any]:
        """导出伏笔表 (EXPORT-006)

        Args:
            project_id: 项目ID
            book_id: 书籍ID
            format: 导出格式 (md/txt)
            status_filter: 状态过滤 (open/progressing/resolved/deferred)

        Returns:
            导出记录
        """
        # 获取伏笔数据
        foreshadows = self._get_foreshadowing_data(book_id, status_filter)

        if not foreshadows:
            raise ValueError("没有伏笔数据可导出")

        # 生成输出文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        book = self.db.fetchone("SELECT title FROM books WHERE id=?", (book_id,))
        book_title = book["title"] if book else "unknown"
        filename = f"{_sanitize_filename(book_title)}_foreshadowing_{timestamp}.{format}"
        file_path = self.output_dir / filename

        # 根据格式导出
        if format == "md":
            content = self._export_foreshadowing_markdown(foreshadows)
        elif format == "txt":
            content = self._export_foreshadowing_txt(foreshadows)
        else:
            raise ValueError(f"不支持的格式: {format}")

        # 写入文件
        file_path.write_text(content, encoding="utf-8")

        # 记录导出
        export_id = generate_id()
        now = datetime.now().isoformat()
        with self.db.transaction() as conn:
            conn.execute(
                """INSERT INTO exports(id, project_id, book_id, format, file_path,
                   file_size, status, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (export_id, project_id, book_id, format, str(file_path),
                 file_path.stat().st_size, "completed", now),
            )

        return {
            "id": export_id,
            "file_path": str(file_path),
            "format": format,
            "file_size": file_path.stat().st_size,
        }

    def _get_foreshadowing_data(
        self,
        book_id: str,
        status_filter: Optional[str] = None,
    ) -> list[dict[str, Any]]:
        """获取伏笔数据"""
        if status_filter:
            foreshadows = self.db.fetchall(
                """SELECT * FROM foreshadows
                   WHERE book_id = ? AND status = ?
                   ORDER BY created_chapter""",
                (book_id, status_filter),
            )
        else:
            foreshadows = self.db.fetchall(
                """SELECT * FROM foreshadows
                   WHERE book_id = ?
                   ORDER BY created_chapter""",
                (book_id,),
            )

        return [dict(f) for f in foreshadows]

    def _export_foreshadowing_markdown(self, foreshadows: list[dict]) -> str:
        """导出伏笔表为 Markdown"""
        lines = [
            "# 伏笔表",
            "",
            "---",
            "",
        ]

        # 统计
        open_count = sum(1 for f in foreshadows if f.get("status") == "open")
        resolved_count = sum(1 for f in foreshadows if f.get("status") == "resolved")
        lines.append(f"**总数**: {len(foreshadows)}")
        lines.append(f"**未解决**: {open_count}")
        lines.append(f"**已解决**: {resolved_count}")
        lines.append("")
        lines.append("---")
        lines.append("")

        # 表格
        lines.append("| ID | 标题 | 描述 | 埋设章节 | 解决章节 | 状态 |")
        lines.append("|---|------|------|---------|---------|------|")

        for f in foreshadows:
            fid = f.get("id", "")[:8]
            title = f.get("title", "")
            description = f.get("description", "")
            created = f.get("created_chapter", "")
            resolved = f.get("resolved_chapter", "")
            status = f.get("status", "")

            lines.append(f"| {fid} | {title} | {description} | {created} | {resolved} | {status} |")

        lines.append("")

        return "\n".join(lines)

    def _export_foreshadowing_txt(self, foreshadows: list[dict]) -> str:
        """导出伏笔表为纯文本"""
        lines = [
            "伏笔表",
            "=" * 20,
            "",
        ]

        # 统计
        open_count = sum(1 for f in foreshadows if f.get("status") == "open")
        resolved_count = sum(1 for f in foreshadows if f.get("status") == "resolved")
        lines.append(f"总数: {len(foreshadows)}")
        lines.append(f"未解决: {open_count}")
        lines.append(f"已解决: {resolved_count}")
        lines.append("")
        lines.append("-" * 40)
        lines.append("")

        for f in foreshadows:
            fid = f.get("id", "")[:8]
            title = f.get("title", "")
            description = f.get("description", "")
            created = f.get("created_chapter", "")
            resolved = f.get("resolved_chapter", "")
            status = f.get("status", "")

            lines.append(f"[{fid}] {title}")
            lines.append(f"  {description}")
            lines.append(f"  埋设: 第{created}章 | 解决: 第{resolved}章 | 状态: {status}")
            lines.append("")

        return "\n".join(lines)
