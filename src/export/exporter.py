"""导出系统 - 支持 docx、md、txt 格式"""

from pathlib import Path
from datetime import datetime
from typing import Optional

from ..core.models import StoryProject


class Exporter:
    """统一导出器"""

    def __init__(self, output_dir: str = "exports"):
        self.output_dir = Path(output_dir)

    def export(self, project: StoryProject, format: str = "md",
               output_path: Optional[str] = None, approved_only: bool = False) -> str:
        """导出小说

        Args:
            project: 项目对象
            format: 导出格式 (md/txt/docx)
            output_path: 自定义输出路径
            approved_only: 只导出已通过审查的章节

        Returns:
            导出文件路径
        """
        self.output_dir.mkdir(parents=True, exist_ok=True)

        if output_path:
            base_path = Path(output_path)
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_path = self.output_dir / f"{project.name}_{timestamp}"

        if format == "md":
            return self._export_markdown(project, base_path.with_suffix(".md"), approved_only)
        elif format == "txt":
            return self._export_txt(project, base_path.with_suffix(".txt"), approved_only)
        elif format == "docx":
            return self._export_docx(project, base_path.with_suffix(".docx"), approved_only)
        else:
            raise ValueError(f"不支持的格式: {format}")

    def _get_sorted_chapters(self, project: StoryProject, approved_only: bool) -> list:
        """获取排序后的章节列表"""
        chapters = list(project.chapters.values())
        if approved_only:
            chapters = [ch for ch in chapters if ch.status.value in ("approved", "exported")]
        return sorted(chapters, key=lambda ch: ch.number)

    def _export_markdown(self, project: StoryProject, path: Path,
                         approved_only: bool) -> str:
        """导出为Markdown"""
        chapters = self._get_sorted_chapters(project, approved_only)

        lines = []
        lines.append(f"# {project.name}\n")
        lines.append(f"类型: {project.genre}\n")
        lines.append(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        lines.append(f"章节数: {len(chapters)}\n")
        lines.append("---\n\n")

        # 目录
        lines.append("## 目录\n\n")
        for ch in chapters:
            lines.append(f"- 第{ch.number}章 {ch.title}\n")
        lines.append("\n---\n\n")

        # 正文
        for ch in chapters:
            lines.append(f"## 第{ch.number}章 {ch.title}\n\n")
            lines.append(ch.content)
            lines.append("\n\n---\n\n")

        content = "".join(lines)
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        return str(path)

    def _export_txt(self, project: StoryProject, path: Path,
                    approved_only: bool) -> str:
        """导出为TXT"""
        chapters = self._get_sorted_chapters(project, approved_only)

        lines = []
        lines.append(f"{project.name}\n")
        lines.append(f"{'=' * 40}\n\n")

        for ch in chapters:
            lines.append(f"第{ch.number}章 {ch.title}\n")
            lines.append(f"{'-' * 30}\n")
            lines.append(ch.content)
            lines.append("\n\n")

        content = "".join(lines)
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        return str(path)

    def _export_docx(self, project: StoryProject, path: Path,
                     approved_only: bool) -> str:
        """导出为DOCX"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        chapters = self._get_sorted_chapters(project, approved_only)
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = getattr(style, 'font')
        font.name = '宋体'
        font.size = Pt(12)

        # 标题
        title = doc.add_heading(project.name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 信息页
        doc.add_paragraph(f"类型: {project.genre}")
        doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"章节数: {len(chapters)}")
        doc.add_page_break()

        # 目录页
        doc.add_heading("目录", level=1)
        for ch in chapters:
            doc.add_paragraph(f"第{ch.number}章 {ch.title}", style='List Number')
        doc.add_page_break()

        # 正文
        for ch in chapters:
            doc.add_heading(f"第{ch.number}章 {ch.title}", level=1)

            # 分段处理正文
            paragraphs = ch.content.split("\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
                    p.paragraph_format.line_spacing = 1.5  # 行距

            doc.add_page_break()

        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))

        return str(path)

    def export_review_report(self, project: StoryProject, format: str = "md") -> str:
        """导出审查报告"""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = self.output_dir / f"{project.name}_审查报告_{timestamp}.md"

        lines = []
        lines.append(f"# {project.name} - 审查报告\n")
        lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

        for num, ch in sorted(project.chapters.items()):
            if ch.review:
                lines.append(f"## 第{num}章: {ch.title}\n")
                lines.append(f"- 总分: {ch.review.overall_score:.1f}\n")
                lines.append(f"- 结论: {ch.review.verdict.value}\n")
                lines.append(f"- 修订次数: {ch.revision_count}\n")

                if ch.review.dimensions:
                    lines.append("\n### 维度评分\n")
                    for dim in ch.review.dimensions:
                        lines.append(f"- {dim.name}: {dim.score:.1f}")
                        if dim.issues:
                            for issue in dim.issues:
                                lines.append(f"  - 问题: {issue}")
                    lines.append("")

                if ch.review.specific_issues:
                    lines.append("\n### 遗留问题\n")
                    for issue in ch.review.specific_issues:
                        lines.append(f"- {issue}")
                    lines.append("")

                lines.append("---\n\n")

        content = "".join(lines)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        return str(path)

    def export_project_documents(self, project: StoryProject) -> dict:
        """导出项目文档集（每卷、每段弧、每章的任务文档）"""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        docs_dir = self.output_dir / f"{project.name}_项目文档"
        docs_dir.mkdir(parents=True, exist_ok=True)

        exported = {}

        # 1. 世界观文档
        world_doc = docs_dir / "世界观设定.md"
        with open(world_doc, "w", encoding="utf-8") as f:
            f.write(f"# {project.name} - 世界观设定\n\n")
            w = project.world
            if w.setting_description:
                f.write(f"## 世界背景\n{w.setting_description}\n\n")
            if w.core_conflict:
                f.write(f"## 核心矛盾\n{w.core_conflict}\n\n")
            if w.power_system:
                f.write(f"## 力量体系\n{w.power_system}\n\n")
            if w.world_rules:
                f.write("## 世界规则\n")
                for rule in w.world_rules:
                    f.write(f"- {rule}\n")
                f.write("\n")
        exported["world"] = str(world_doc)

        # 2. 角色文档
        char_doc = docs_dir / "角色设定.md"
        with open(char_doc, "w", encoding="utf-8") as f:
            f.write(f"# {project.name} - 角色设定\n\n")
            for name, char in project.characters.items():
                f.write(f"## {name}\n")
                f.write(f"- 角色定位: {char.role}\n")
                if char.personality:
                    f.write(f"- 性格: {char.personality}\n")
                if char.background:
                    f.write(f"- 背景: {char.background}\n")
                if char.abilities:
                    f.write(f"- 能力: {', '.join(char.abilities)}\n")
                if char.relationships:
                    f.write("- 关系:\n")
                    for rel_name, rel_type in char.relationships.items():
                        f.write(f"  - {rel_name}: {rel_type}\n")
                f.write("\n")
        exported["characters"] = str(char_doc)

        # 3. 卷规划文档
        for vol in project.volumes:
            vol_doc = docs_dir / f"第{vol.number}卷_{vol.title}.md"
            with open(vol_doc, "w", encoding="utf-8") as f:
                f.write(f"# {vol.title}\n\n")
                f.write(f"{vol.description}\n\n")
                if vol.themes:
                    f.write(f"主题: {', '.join(vol.themes)}\n\n")
                for arc in vol.arcs:
                    f.write(f"## 段弧: {arc.name}\n")
                    f.write(f"{arc.description}\n")
                    if arc.key_events:
                        f.write("关键事件:\n")
                        for event in arc.key_events:
                            f.write(f"- {event}\n")
                    f.write("\n")
            exported[f"vol_{vol.number}"] = str(vol_doc)

        # 4. 伏笔文档
        hook_doc = docs_dir / "伏笔清单.md"
        with open(hook_doc, "w", encoding="utf-8") as f:
            f.write(f"# {project.name} - 伏笔清单\n\n")
            for fid, hook in project.foreshadowing.items():
                f.write(f"## [{fid}] {hook.description}\n")
                f.write(f"- 状态: {hook.status}\n")
                f.write(f"- 埋设章节: 第{hook.planted_chapter}章\n")
                if hook.resolved_chapter:
                    f.write(f"- 回收章节: 第{hook.resolved_chapter}章\n")
                if hook.related_characters:
                    f.write(f"- 相关人物: {', '.join(hook.related_characters)}\n")
                f.write("\n")
        exported["foreshadowing"] = str(hook_doc)

        return exported
